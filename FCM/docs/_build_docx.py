"""Generate FCM_Methodology.docx — the formal methodology write-up that
accompanies the Gulf FEI Fuzzy Cognitive Map subsystem. Re-run any time the
methodology evolves.

    python -m FCM.docs._build_docx
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

OUT_PATH = Path(__file__).resolve().parent / "FCM_Methodology.docx"

NAVY = RGBColor(0x01, 0x30, 0x56)
TEAL = RGBColor(0x05, 0x6F, 0x73)
GREY = RGBColor(0x55, 0x55, 0x55)


# ── Style helpers ──────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 1 else TEAL
        run.font.name = "Calibri"


def _add_para(doc: Document, text: str, *, italic: bool = False,
              bold: bool = False, size: int = 11, color: RGBColor = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.italic = italic
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def _add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = "Calibri"


def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = GREY


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10.5)
        run.font.name = "Calibri"
        _set_cell_shading(cell, "013056")
    for r, row_vals in enumerate(rows, start=1):
        for c, val in enumerate(row_vals):
            cell = table.rows[r].cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.name = "Calibri"


# ── Document builder ───────────────────────────────────────────────────────

def build() -> Document:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)

    # ── Title page ─────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Gulf FEI Fuzzy Cognitive Map (FCM) Subsystem")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Methodology, Architecture, and Quality Controls")
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = TEAL
    run.font.name = "Calibri"

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"Prepared for the Gulf FEI research team\n"
        f"Gulf of Mexico Fishery Ecosystem Initiative\n"
        f"Document version: {date.today().isoformat()}"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    run.font.name = "Calibri"

    doc.add_page_break()

    # ── 1. Executive summary ────────────────────────────────────────────
    _add_heading(doc, "1. Executive Summary", 1)
    _add_para(doc,
        "This document describes how the Gulf FEI platform turns scientific "
        "documents and user research questions into Fuzzy Cognitive Maps "
        "(FCMs) — signed, weighted, directed graphs that summarise the "
        "causal structure of the Gulf of Mexico fisheries ecosystem in a "
        "form suitable for ‘what-if’ scenario analysis and policy "
        "discussion. It covers the theoretical motivation, the end-to-end "
        "pipeline, the quality controls that keep the keyword set honest, "
        "and the limitations the team should be aware of when sharing "
        "outputs with external stakeholders."
    )

    # ── 2. Background ───────────────────────────────────────────────────
    _add_heading(doc, "2. Background and Motivation", 1)
    _add_para(doc,
        "Fuzzy Cognitive Maps were introduced by Bart Kosko (1986) as a "
        "way to combine the interpretability of a concept graph with the "
        "dynamics of a recurrent neural network. Each node is a domain "
        "concept (e.g. “Overfishing”, “Coral Bleaching”) and each "
        "directed edge carries a signed weight in [−1, +1] representing the "
        "strength and direction of causal influence. Once the matrix is "
        "assembled, activation can be propagated iteratively until the "
        "system reaches equilibrium, allowing analysts to ask questions "
        "such as “if fishing pressure stays elevated and habitat continues "
        "to degrade, which concepts settle into a depressed state?”"
    )
    _add_para(doc,
        "FCMs are an excellent fit for NOAA’s Ecosystem-Based Fisheries "
        "Management (EBFM) framework because they:")
    _add_bullets(doc, [
        "Bridge qualitative expert knowledge and quantitative simulation "
        "without requiring full mechanistic models.",
        "Are auditable — every edge in the Gulf FEI implementation carries "
        "the source sentence as evidence so reviewers can verify the claim.",
        "Support participatory science — stakeholders can amend the graph "
        "by adding manual edges or uploading their own adjacency matrices.",
        "Scale across data sources — the same graph can be built from RAG "
        "retrieval, uploaded PDFs/Word docs, or pre-existing CSV/XLSX matrices.",
    ])

    # ── 3. System architecture ──────────────────────────────────────────
    _add_heading(doc, "3. System Architecture", 1)
    _add_para(doc,
        "The FCM subsystem is one stage of a wider Retrieval-Augmented "
        "Generation pipeline. The high-level flow is:"
    )
    _add_code_block(doc,
        "User question / uploaded file\n"
        "        |\n"
        "        v\n"
        "  RAG retrieval + LLM answer        (src/rag.py, src/vector_loader.py)\n"
        "        |\n"
        "        v\n"
        "  CausalExtractor                   (FCM/extractor.py)\n"
        "        |  concepts + signed edges\n"
        "        v\n"
        "  FCMGraphBuilder                   (FCM/fcm_graph.py)\n"
        "        |  FCMMap (concepts, edges, adjacency_matrix)\n"
        "        v\n"
        "  FCMGenerator                      (src/FCM.py)\n"
        "        |  networkx graph + publication-quality plots\n"
        "        v\n"
        "  FCMSimulator                      (FCM/simulator.py)\n"
        "        |  scenario equilibria\n"
        "        v\n"
        "  FastAPI response (JSON + PNGs)    (main.py)\n"
    )
    _add_heading(doc, "3.1 Module responsibilities", 2)
    _add_table(doc,
        ["Module", "Responsibility"],
        [
            ["FCM/schemas.py", "Pydantic data contracts (Edge, FCMMap, request/response models)."],
            ["FCM/extractor.py", "Causal extraction: regex patterns, co-occurrence, transitive 2-hop, fuzzy merge, polarity interleave."],
            ["FCM/fcm_graph.py", "Assembles the canonical FCMMap and exports JSON / GEXF / CSV / PNG artefacts."],
            ["FCM/categorizer.py", "FEP colour scheme + summary-theme rules (mirrors the original R pipeline)."],
            ["FCM/simulator.py", "Kosko activation propagation with sigmoid / tanh squashing and driver clamping."],
            ["src/FCM.py", "High-level orchestrator + publication renderer used by the FastAPI app."],
            ["main.py", "FastAPI endpoints (/query, /upload-files, /simulate-fcm) and request validation."],
        ],
    )

    # ── 4. Pipeline stages ──────────────────────────────────────────────
    _add_heading(doc, "4. Pipeline Stages in Detail", 1)

    _add_heading(doc, "4.1 Document ingestion and RAG retrieval", 2)
    _add_para(doc,
        "Two entry points feed the pipeline. The /query endpoint accepts a "
        "free-text research question, retrieves the top-k most relevant "
        "chunks from the FAISS vector store, and constructs a single text "
        "bundle of the form Q + A + Context. The /upload-files endpoint "
        "accepts up to 20 PDF / DOCX / TXT / CSV / XLSX files; any uploaded "
        "adjacency matrix bypasses extraction entirely and is consolidated "
        "directly via master-node expansion (mirroring the R pipeline)."
    )

    _add_heading(doc, "4.2 Causal extraction", 2)
    _add_para(doc,
        "The extractor combines three complementary strategies so the graph "
        "remains useful even when the source text uses indirect phrasing:"
    )
    _add_numbered(doc, [
        "Pattern matching — a registry of compiled regex templates "
        "(CAUSAL_PATTERNS) detects explicit causal verbs ('increases', "
        "'reduces', 'leads to', 'due to', 'is associated with', …) and "
        "assigns a base weight in [−0.75, +0.75] and a polarity label.",
        "Co-occurrence fallback — when explicit patterns are sparse, any "
        "two domain keywords (matched against _DOMAIN_KW) appearing in the "
        "same 3-sentence rolling window are connected with a "
        "distance-weighted edge. The window's dominant sentiment "
        "(_NEG_MARKERS vs _POS_MARKERS) sets the sign so co-occurrence "
        "edges are not blindly positive.",
        "Transitive closure — a damped 2-hop inference fires "
        "A→B + B→C ⇒ A→C with weight = 0.7·w_AB·w_BC. These edges "
        "are tagged edge_type='transitive', hops=2 so reviewers can filter "
        "or visually distinguish them from primary, evidence-backed edges.",
    ])
    _add_para(doc,
        "All three strategies write into a single (concept, edge) pool that "
        "is then deduplicated and ranked."
    )

    _add_heading(doc, "4.3 Concept normalisation and cleanliness", 2)
    _add_para(doc,
        "Every concept name passes through two strict gates before it can "
        "appear in the final graph. These gates were hardened in the latest "
        "release after the team observed clause fragments and useless "
        "modifiers leaking into the keyword list."
    )
    _add_para(doc, "normalize_concept performs:", bold=True)
    _add_bullets(doc, [
        "Strips leading articles (the / a / an).",
        "Strips trailing prepositions / conjunctions / articles iteratively.",
        "Strips leading and trailing useless modifiers ('very', 'high', "
        "'much', 'good', …).",
        "Collapses consecutive duplicate tokens ('Fish Fish Stock' → 'Fish Stock').",
        "Title-cases the result while preserving domain acronyms (CPUE, MSY, "
        "NOAA, EBFM …).",
    ])
    _add_para(doc, "_is_clean_concept rejects a candidate if any of:", bold=True)
    _add_bullets(doc, [
        "It is empty, longer than 4 words, or shorter than 4 characters.",
        "The first word is a causal verb, gerund, past tense, pronoun, "
        "preposition, or interrogative.",
        "The last word is a dangling preposition or conjunction.",
        "It is a single-word past-participle adjective ('increased', "
        "'reducing') and not in the domain whitelist (so 'Overfishing', "
        "'Spawning', 'Fishing' are kept).",
        "It is a single-word vague noun ('thing', 'topic', 'level', …).",
        "Any word in the span is an interrogative, auxiliary, pronoun, or "
        "negation — the span is then a clause fragment, not a noun phrase.",
        "Every word is a useless modifier ('Very High', 'Much More').",
        "More than half of the words are stopwords.",
        "Any token is shorter than 2 characters or the words are all duplicates.",
        "It is a pure-numeric span ('2', '100', '3.5').",
        "It contains a clear passive/relational verb marker ('associated', "
        "'linked', 'correlated', …).",
    ])

    _add_heading(doc, "4.4 Fuzzy merging and bidirectional dedup", 2)
    _add_para(doc,
        "Once cleaned, near-duplicate concepts are merged with a two-stage "
        "process: a lemma key (singular / sorted-token) collects obvious "
        "duplicates, then a SequenceMatcher pass (default similarity ≥ 0.88) "
        "collapses fuzzy variants. The longer name wins as canonical because "
        "it tends to be more informative. Edges between the same "
        "(source, target) post-merge are combined by confidence-weighted mean."
    )
    _add_para(doc,
        "_dedupe_bidirectional then drops same-sign reverse pairs (A→B and "
        "B→A both positive) keeping the stronger direction. Opposing-sign "
        "pairs are preserved because they encode legitimate FCM feedback "
        "loops."
    )

    _add_heading(doc, "4.5 Polarity interleaving", 2)
    _add_para(doc,
        "A credible signed FCM should not be dominated by one polarity in "
        "the first N edges shown to the reviewer. _interleave_polarity "
        "preserves the overall ranking within each polarity but interleaves "
        "them at a 2:1 cadence (positive-heavy) so negative findings remain "
        "visible without overwhelming the typical positive signal."
    )

    _add_heading(doc, "4.6 Adjacency matrix and exports", 2)
    _add_para(doc,
        "FCMGraphBuilder packages the cleaned concept list and edges into "
        "the canonical FCMMap (concepts in alphabetical order, adjacency "
        "matrix indexed by that order). export_network writes JSON, GEXF, "
        "two CSVs (adjacency + edge list), per-node metric CSV, and a "
        "publication-quality PNG to disk."
    )

    _add_heading(doc, "4.7 Visualisation", 2)
    _add_para(doc,
        "FCMGenerator.plot_fcm renders the FCM with several research-grade "
        "design choices:"
    )
    _add_bullets(doc, [
        "Layout — Kamada–Kawai energy minimisation with a hard "
        "minimum-distance floor so nodes never overlap.",
        "Node size — proportional to a 0.55 * betweenness + 0.45 * degree "
        "centrality blend (importance → visual weight).",
        "Node colour — community detected by greedy modularity, drawn from "
        "the colour-blind-safe Okabe-Ito palette.",
        "Role ring colour — transmitter (orange), receiver (blue), ordinary "
        "(green), isolated (grey), classified by in/out strength ratios.",
        "Edges — green for positive, red for negative; transitive 2-hop "
        "edges drawn dashed and faded so reviewers can tell them apart from "
        "evidence-backed primary edges.",
        "Embedded statistics panel — node count, edge count, density, "
        "average degree, role distribution, top-5 concepts by total "
        "strength, and module count.",
    ])
    _add_para(doc,
        "A second 'summary thematic' rendering (plot_summary_fcm) collapses "
        "fine-grained nodes into the categorizer's super-themes and renders "
        "the same graph at a policy-relevant level of abstraction."
    )

    _add_heading(doc, "4.8 Scenario simulation", 2)
    _add_para(doc,
        "FCMSimulator implements Kosko's propagation rule "
        "a(t+1) = f( a(t) + A^T a(t) ) where A is the signed adjacency "
        "matrix, a(t) is the activation vector, and f is a squashing "
        "function (sigmoid, tanh, or identity). User-selected drivers are "
        "clamped to their initial value at every step so the resulting "
        "equilibrium answers a clean 'what if X stays elevated' question. "
        "Convergence is detected when max |a(t+1) − a(t)| falls below 1e-4 "
        "or after the configured iteration cap."
    )

    # ── 5. Quality controls table ───────────────────────────────────────
    _add_heading(doc, "5. Quality Controls and Validation", 1)
    _add_para(doc,
        "Several guards keep the keyword set and graph honest. Each one is "
        "defensive in a different way — together they ensure that what "
        "appears in the final FCM is a defensible noun phrase backed by "
        "evidence the reviewer can audit."
    )
    _add_table(doc,
        ["Guard", "Where", "What it does"],
        [
            ["Query meaningfulness check",
             "main._is_meaningful_query",
             "Rejects empty / pure-stopword / interrogative-only queries before the RAG + FCM pipeline runs."],
            ["Concept normalisation",
             "extractor.normalize_concept",
             "Strips articles, leading/trailing useless modifiers, dedupes repeated tokens, preserves acronyms."],
            ["Concept cleanliness",
             "extractor._is_clean_concept",
             "Rejects clause fragments, interrogatives anywhere, all-modifier phrases, sub-2-letter tokens, vague singles, >50% stopword spans."],
            ["Fuzzy merging",
             "extractor.fuzzy_merge_concepts",
             "Collapses near-duplicate concepts (plurals, fuzzy matches) into the longest canonical name; merges edges by confidence-weighted mean."],
            ["Bidirectional dedup",
             "extractor._dedupe_bidirectional",
             "Drops same-sign reverse pairs keeping the stronger direction; preserves opposing-sign pairs as legitimate feedback loops."],
            ["Transitive closure damping",
             "extractor._transitive_closure",
             "Damps inferred 2-hop edges by 0.7 and tags them edge_type='transitive' so the UI can visually distinguish them."],
            ["Polarity interleaving",
             "extractor._interleave_polarity",
             "Re-ranks final edges so positive and negative findings stay visible in the top-N (2:1 cadence)."],
            ["Polarity warnings",
             "main._build_response",
             "Flags maps that contain only positive or only negative edges so reviewers can sanity-check the source text."],
        ],
    )

    # ── 6. Inputs / outputs ─────────────────────────────────────────────
    _add_heading(doc, "6. Inputs and Outputs", 1)
    _add_heading(doc, "6.1 Inputs accepted by the public API", 2)
    _add_bullets(doc, [
        "POST /query with form fields (query, max_edges) — standard RAG "
        "+ FCM path used by the web UI.",
        "POST /api/query with JSON {question, relation_length} — same "
        "pipeline for programmatic use.",
        "POST /upload-files with one or more files (.pdf, .docx, .txt, .md, "
        ".csv, .xlsx) — if any uploaded CSV/XLSX matches the adjacency "
        "matrix layout it is consolidated directly; otherwise all text is "
        "combined and run through extraction.",
        "POST /simulate-fcm with JSON {activations, steps, squash, lam, "
        "clamp_drivers} — runs Kosko propagation on the most recently "
        "built FCM.",
    ])
    _add_heading(doc, "6.2 Outputs", 2)
    _add_bullets(doc, [
        "answer — the LLM's natural-language response to the query.",
        "fcm_edge_results — ranked list of (cause, effect, polarity, "
        "weight) tuples ready for tabular display.",
        "edges_meta — same edges enriched with confidence, edge_type, "
        "hops, evidence sentence, and evidence document id.",
        "adjacency_matrix + matrix_concepts — the signed matrix.",
        "plot_path / plot_path_summary — PNG locations for the "
        "consolidated and summary-thematic renderings.",
        "warnings — sanity-check messages (e.g. 'no negative edges').",
        "Simulation response — final activation, per-concept influence "
        "(Δ from baseline), iteration count, and convergence flag.",
    ])

    # ── 7. Use cases ────────────────────────────────────────────────────
    _add_heading(doc, "7. Use Cases for the Research Team", 1)
    _add_numbered(doc, [
        "Literature triage — ask a research question and read off the "
        "ranked causal claims, then click through to the evidence sentence "
        "for any edge that surprises you.",
        "Stakeholder workshops — export the consolidated FCM as PNG and "
        "the adjacency matrix as CSV, share with subject-matter experts, "
        "ask them to amend weights or add manual edges.",
        "Cross-document synthesis — upload a set of NOAA reports and "
        "SEDAR assessments together; the consolidator merges their adjacency "
        "matrices into a single master graph.",
        "Scenario analysis — clamp Overfishing high and Habitat "
        "Restoration low, run the simulator, observe which downstream "
        "concepts collapse versus stabilise.",
        "Pipeline auditing — every edge carries its source sentence and "
        "an edge_type label so reviewers can distinguish high-confidence "
        "regex-pattern edges from co-occurrence and inferred edges.",
    ])

    # ── 8. Limitations ──────────────────────────────────────────────────
    _add_heading(doc, "8. Limitations and Future Work", 1)
    _add_bullets(doc, [
        "Pattern coverage — the regex registry handles the common causal "
        "verb constructions but will miss subtle phrasing. Expanding "
        "CAUSAL_PATTERNS is a low-cost ongoing improvement.",
        "Single-language — the extractor is English-only. Adding a "
        "Spanish lexicon would extend coverage to additional Gulf "
        "stakeholders.",
        "Weight calibration — base weights for each pattern are heuristic. "
        "Future work could fit them to expert-rated reference graphs.",
        "Domain whitelist drift — _DOMAIN_KW is a fishery-focused regex; "
        "extending the FCM to adjacent domains (ocean policy, aquaculture "
        "supply chain) would require curating a new whitelist.",
        "Simulation interpretation — Kosko equilibria are qualitative "
        "indicators, not predictions. Outputs should always be presented "
        "alongside the underlying evidence and discussed with domain experts.",
    ])

    # ── 9. References ───────────────────────────────────────────────────
    _add_heading(doc, "9. References and Reproducibility", 1)
    _add_bullets(doc, [
        "Kosko, B. (1986). Fuzzy Cognitive Maps. International Journal of "
        "Man-Machine Studies, 24(1), 65–75.",
        "Papageorgiou, E. I., & Salmeron, J. L. (2013). A review of fuzzy "
        "cognitive maps research during the last decade. IEEE Transactions "
        "on Fuzzy Systems, 21(1), 66–79.",
        "NOAA Ecosystem-Based Fisheries Management Roadmap. National Marine "
        "Fisheries Service.",
        "Companion notebook — FCM/docs/FCM_Walkthrough.ipynb (run "
        "cell-by-cell to reproduce every artefact described here).",
        "Source code — the FCM/ package in this repository; entry points "
        "are documented in main.py.",
    ])

    return doc


def main() -> None:
    doc = build()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
